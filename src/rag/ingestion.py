import json
import sys
from pathlib import Path
from typing import List
 
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
 
# Agregar raíz del proyecto al path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from config import CHUNK_SIZE, CHUNK_OVERLAP, CVS_DIR, EVALUACIONES_DIR
 
 
# Loaders por tipo de archivo 
 
def _load_pdf(path: str, candidate_id: str, source_type: str) -> List[Document]:
    loader = PyPDFLoader(path)
    docs = loader.load()
    for d in docs:
        d.metadata.update({"candidate_id": candidate_id, "source_type": source_type})
    return docs
 
 
def _load_txt(path: str, candidate_id: str, source_type: str) -> List[Document]:
    loader = TextLoader(path, encoding="utf-8")
    docs = loader.load()
    for d in docs:
        d.metadata.update({"candidate_id": candidate_id, "source_type": source_type})
    return docs
 
 
def _load_json(path: str, candidate_id: str, source_type: str) -> List[Document]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    content = json.dumps(data, ensure_ascii=False, indent=2)
    return [Document(
        page_content=content,
        metadata={"candidate_id": candidate_id, "source_type": source_type, "source": path},
    )]
 
 
def _load_docx(path: str, candidate_id: str, source_type: str) -> List[Document]:
    try:
        import docx
        doc = docx.Document(path)
        content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(
            page_content=content,
            metadata={"candidate_id": candidate_id, "source_type": source_type, "source": path},
        )]
    except ImportError:
        print("[ingestion] ⚠ python-docx no instalado. Instala con: pip install python-docx")
        return []
 
 
def _infer_source_type(filename: str) -> str:
    name = filename.lower()
    if any(k in name for k in ["cv", "curricul", "resume", "hoja_de_vida"]):
        return "curriculum"
    if "linkedin" in name:
        return "linkedin"
    if "github" in name:
        return "github"
    if any(k in name for k in ["feedback", "entrevista", "interview"]):
        return "feedback_entrevista"
    if any(k in name for k in ["evaluacion", "eval", "assessment"]):
        return "evaluacion_previa"
    return "otro"
 
 
# Cargador desde archivo único (para subida desde UI)
 
def load_single_file(
    file_path: str,
    candidate_id: str,
    source_type: str = None,
) -> List[Document]:
    """
    Carga un único archivo y retorna sus Documents.
    Usado cuando el usuario sube archivos desde la interfaz Streamlit.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    inferred_type = source_type or _infer_source_type(path.name)
 
    loaders = {
        ".pdf":  lambda: _load_pdf(file_path, candidate_id, inferred_type),
        ".txt":  lambda: _load_txt(file_path, candidate_id, inferred_type),
        ".json": lambda: _load_json(file_path, candidate_id, inferred_type),
        ".docx": lambda: _load_docx(file_path, candidate_id, inferred_type),
    }
 
    loader_fn = loaders.get(ext)
    if not loader_fn:
        print(f"[ingestion] Extensión no soportada: {ext}")
        return []
 
    docs = loader_fn()
    print(f"[ingestion] ✓ {path.name} ({inferred_type}) → {len(docs)} doc(s)")
    return docs
 
 
# Cargador desde directorio completo
 
def load_all_candidates() -> List[Document]:
    """
    Carga todos los candidatos desde data/cvs/ y data/evaluaciones/.
    Estructura esperada:
        data/cvs/<candidate_id>/archivo.pdf
        data/evaluaciones/<candidate_id>/feedback.txt
    """
    all_docs: List[Document] = []
 
    for base_dir in [CVS_DIR, EVALUACIONES_DIR]:
        for candidate_folder in sorted(base_dir.iterdir()):
            if not candidate_folder.is_dir():
                continue
            candidate_id = candidate_folder.name
 
            for file in candidate_folder.iterdir():
                ext = file.suffix.lower()
                source_type = _infer_source_type(file.name)
                docs = load_single_file(str(file), candidate_id, source_type)
                all_docs.extend(docs)
 
    print(f"[ingestion] Total documentos cargados: {len(all_docs)}")
    return all_docs
 
 
# Chunking
 
def split_documents(documents: List[Document]) -> List[Document]:
    """Divide documentos largos en fragmentos para embeddings."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    print(f"[ingestion] Chunks generados: {len(chunks)}")
    return chunks